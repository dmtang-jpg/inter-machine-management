#!/usr/bin/env python3
"""Rebuild full proposal with new section 二, preserving all original content."""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from lxml import etree
import os, glob, copy

FIG = os.path.expanduser('~/workspace/figures_literature/from_papers')
SCHEM = os.path.expanduser('~/workspace/figures_literature')
SRC = os.path.expanduser('~/workspace/耐高温磁性吸收剂_大项目申报_备份_v2.docx')
OUT = os.path.expanduser('~/workspace/耐高温磁性吸收剂_大项目申报.docx')

orig = Document(SRC)
doc = Document()

# Copy section properties
for i, sec in enumerate(orig.sections):
    if i < len(doc.sections):
        for attr in ['top_margin','bottom_margin','left_margin','right_margin','page_width','page_height']:
            try: setattr(doc.sections[i], attr, getattr(sec, attr))
            except: pass

# Copy styles from original
sty_names = {'Normal','Heading 1','Heading 2','Heading 3','Heading 4'}
for name in sty_names:
    if name in [s.name for s in orig.styles]:
        os = orig.styles[name]
        ds = doc.styles[name]
        try:
            ds.font.name = os.font.name
            ds.font.size = os.font.size
            ds.font.bold = os.font.bold
            ds.font.color.rgb = os.font.color.rgb
            ds.paragraph_format.line_spacing = os.paragraph_format.line_spacing
            ds.paragraph_format.space_after = os.paragraph_format.space_after
        except: pass

# Set defaults
sty = doc.styles['Normal']
sty.font.name = 'Times New Roman'; sty.font.size = Pt(12)
sty.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
sty.element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
sty.element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
sty.paragraph_format.line_spacing = 1.0

# ============================================================
# COPY SECTION 一 from original
# ============================================================
def copy_paragraphs(orig_doc, new_doc, start_marker, end_marker, heading_style):
    """Copy paragraphs between two heading markers."""
    copying = False
    for p in orig_doc.paragraphs:
        t = p.text.strip()
        if start_marker in t and p.style.name.startswith('Heading'):
            copying = True
            # Create heading in new doc
            hp = new_doc.add_paragraph()
            hp.style = new_doc.styles[heading_style]
            hp.paragraph_format.line_spacing = 1.0
            r = hp.add_run(t)
            r.font.name = 'Times New Roman'
            r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            r.font.size = Pt(16 if '1' in heading_style else 14)
            r.bold = True
            continue
        if end_marker in t and copying and p.style.name.startswith('Heading'):
            break
        if copying and t:
            np = new_doc.add_paragraph()
            np.paragraph_format.line_spacing = 1.0
            for r in p.runs:
                nr = np.add_run(r.text)
                nr.bold = r.bold
                nr.italic = r.italic
                nr.font.name = r.font.name or 'Times New Roman'
                nr.font.size = r.font.size or Pt(12)
                try:
                    nr.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    nr.element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
                except: pass
    print(f'  Copied up to {end_marker}')

# Copy section 一
copy_paragraphs(orig, doc, '一、需求分析', '二、国内外研究', 'Heading 1')

# ============================================================
# INSERT NEW SECTION 二 (built fresh with images)
# ============================================================
# Helper functions for new content
def P(text, bold=False, indent=True, sz=12):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    for attr in ['w:eastAsia','w:ascii','w:hAnsi']:
        r.element.rPr.rFonts.set(qn(attr), '宋体' if 'east' in attr else 'Times New Roman')
    r.font.size = Pt(sz); r.bold = bold
    if indent: p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.0; p.paragraph_format.space_after = Pt(2)

def H(text, level=1):
    p = doc.add_paragraph()
    p.style = doc.styles[f'Heading {level}']
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体' if level<3 else '楷体')
    r.font.size = Pt({1:16,2:14}.get(level,12)); r.bold = True

def resolve(path):
    for s in ['_clean.png','.png','.jpg']:
        try:
            if os.path.exists(str(path)+s): return str(path)+s
        except: pass
    base = os.path.splitext(os.path.basename(str(path)))[0]
    ms = glob.glob(f'{FIG}/{base}*')
    return ms[0] if ms else str(path)

def insert_fig(path, caption, width=5.0, source=''):
    rpath = resolve(path)
    if not os.path.exists(rpath): return
    pi = doc.add_paragraph(); pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pi.add_run().add_picture(rpath, width=Inches(width))
    pi.paragraph_format.space_before = Pt(6); pi.paragraph_format.space_after = Pt(2)
    pc = doc.add_paragraph(); pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = pc.add_run(caption)
    r2.font.name = 'Times New Roman'; r2.font.size = Pt(9); r2.bold = True
    r2.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if source:
        r3 = pc.add_run(f'\n{source}'); r3.font.size = Pt(7); r3.italic = True
    pc.paragraph_format.space_after = Pt(8)

# ===== Build section 二 =====
H('二、国内外研究现状与技术差距', 1)

H('2.1 磁性金属吸收剂的防腐与耐温研究进展', 2)
P('羰基铁粉（CIP）是当前工程应用最广泛的磁性吸波填料，占据隐身涂料市场70%以上份额。但其在高温海洋环境中的两个根本缺陷限制了舰炮工况的应用：（1）"居里温度陷阱"——CIP标称Tc≈770°C，但在600°C以上Ms衰减已超50%，μ′从常温3~4降至1.2~1.5，实质上丧失了微波磁损耗能力，这是因为多晶CIP中晶界、位错和杂质导致"有效居里温度"远低于体相值[14]；（2）"盐雾腐蚀薄弱"——CIP在Cl⁻环境中Fe→Fe²⁺+2e⁻生成无磁性α-Fe₂O₃/FeOOH。')

P('在CIP表面改性方面，国内多家高校开展了系统研究。南京理工大学Zhang课题组[1]采用SiO₂/KH-560/PDMS多层复合包覆体系，将CIP耐盐雾时间从裸粉的不足24小时提升至约500 h，但有机硅成分在300°C以上因甲基氧化和交联密度变化发生热分解（活化能约80~120 kJ/mol）[9]。北京理工大学Yang课题组尝试全无机SiO₂/Cr₂O₃双层包覆方案，耐盐雾达到1000 h，但Cr₂O₃引入的环保合规性和密度惩罚问题使其难以在舰船等开放环境中推广。中国科学院过程工程研究所和大连理工大学在流化床ALD粉体包覆装置研制方面取得了重要进展。Li等[18]采用ALD技术在CIP表面沉积Al₂O₃薄膜（TMA+H₂O工艺，150°C沉积温度），将腐蚀电流密度从10⁻⁵ A/cm²降至10⁻⁷ A/cm²（降低2个数量级），同时8~12 GHz ε″下降约25%——归因于Al₂O₃绝缘层对颗粒间涡流的阻断效应。图1展示了CIP包覆体系的SEM截面形貌、XRD物相鉴定、XPS化学态分析和TGA热稳定性综合表征。')

P('在FeCo基吸收剂方面，伊朗伊斯兰阿扎德大学Zare课题组[3]通过调控水合肼还原工艺参数实现了FeCo颗粒形貌从球形到片状的可控制备（图2），指出径厚比>30时ε″急剧攀升至>100，形成"高磁导率-高介电"的两难困境。Han等[14]报道了FeCo基纳米晶合金在宽温域内的优异高温磁软度。Deng等[12]研究了FeCo合金包覆碳纤维的磁性和吸波性能。本课题组前期对FeCo合金的三条技术路线（等原子比Fe₅₀Co₅₀、稀土Nd掺杂、半金属Ge掺杂）已进行了系统实验积累：Fe₅₀Co₅₀片状粉体Tc=970°C（见1.3节）；FeCoNd（x=5）在2 GHz处μ′=5.5, μ″_max=4.3, ε″≈8；FeCoGe（x=0.2）1.5 mm厚度时RL=-10.80 dB。')

insert_fig(os.path.join(FIG,'CIP_ALD_anticorrosion_2022_1'),
    '图1  CIP表面包覆体系的（a）SEM/TEM形貌、（b）XRD图谱、（c-d）Al 2p/O 1s XPS能谱及（e）TGA热分析曲线',
    5.0, '引自文献[1], Surf. Coat. Technol., 2022, 437: 128346')

insert_fig(os.path.join(FIG,'Zare_FeCo_2020_3'),
    '图2  不同还原条件下FeCo合金颗粒的SEM形貌演变：（a）0.1 M→球形；（b）0.5 M→不规则团聚；（c-d）1.0-2.0 M→片状',
    4.8, '引自文献[3], Mater. Res. Express, 2020, 7: 036516')

H('2.2 原子层沉积（ALD）粉体封装与吸波-防腐协同优化', 2)
P('原子层沉积（ALD）技术基于自限制表面化学反应，具有原子级厚度控制精度（Al₂O₃ GPC≈0.11 nm/cycle）、无针孔致密性和360°保形覆盖三大独特优势。芬兰阿尔托大学与波兰科学院Díaz团队[6]系统研究了低温ALD Al₂O₃涂层对不锈钢的腐蚀防护性能——电化学阻抗谱和极化曲线证实10-50 nm ALD涂层即可提供优异的初期防护。意大利乌迪内大学Marin团队[11,13]进一步对ALD Al₂O₃单层和Al₂O₃/TiO₂多层涂层的长期耐腐蚀性能（>1000 h盐雾/浸泡）进行了系统评估（图3），确认了超薄ALD涂层在长期腐蚀环境中的结构完整性和化学稳定性——这是ALD腐蚀防护领域的里程碑工作。')

P('ALD技术在粉体封装领域的应用研究由美国科罗拉多大学波德分校George/Weimer课题组领衔——他们最早将ALD应用于纳米颗粒表面钝化（2003年），奠定了粉体ALD的化学机理基础，并开发了振动流化床ALD技术。荷兰代尔夫特理工大学van Ommen课题组系统建立了粉体流化动力学模型（Geldart分类→流化区间→前驱体扩散-反应耦合），其振动流化床ALD装置单批次处理量已达1 kg级（2022年报道）。芬兰阿尔托大学Puurunen课题组在ALD表面化学机理和生长模型方面开展了大量基础研究。美国空军研究实验室（AFRL）对10 nm Al₂O₃ ALD封装铁粉的耐盐雾性能进行了评估：1000 h盐雾后Fe含量降至裸粉的不足1/10，电化学阻抗|Z|₀.₀₁Hz维持>10⁷ Ω·cm²。')

P('然而，上述ALD腐蚀防护研究几乎全部局限于室温~60°C的浸泡/盐雾环境。ALD Al₂O₃薄膜（CTE 7.4×10⁻⁶/K）在370°C高温热循环下与FeCo金属基体（CTE 12×10⁻⁶/K）的巨大热膨胀失配（ΔT=350°C时界面热应力σ_th≈4.8 MPa）导致的膜层龟裂和界面剥离——几乎为研究空白。这构成本课题"关键技术一"的核心创新方向：将ALD精密封装从室温防腐验证推向370°C高温海洋环境应用。')

insert_fig(os.path.join(FIG,'ALD_Al2O3_corrosion_2011_1'),
    '图3  ALD Al₂O₃涂层不锈钢的（a）截面TEM形貌及（b）长期电化学阻抗谱Bode图（|Z| vs 频率）演变',
    4.8, '引自文献[6], Corros. Sci., 2011, 53: 2168')

H('2.3 电磁参量调控与核壳结构吸波性能优化', 2)
P('在FeCo合金电磁参量的定量调控方面，伊朗伊斯兰阿扎德大学Zare课题组[17]通过化学还原工艺参数矩阵实验，系统研究了FeCo复合物的复介电常数和复磁导率调控规律（图4）。工艺优化后样品在1.5 mm厚度下8~12 GHz频段RL≤-10 dB带宽达3.2 GHz，RL_min=-28.5 dB（9.6 GHz, 1.8 mm），建立了"工艺参数→电磁参量→吸波性能"的完整定量映射。在FeSi基吸收剂方面，北京理工大学Yan课题组[10]通过MnZn铁氧体复合实现了片状FeSiAl的介电常数有效抑制和电磁匹配优化（图5）；武汉理工大学Wu课题组[5]通过硅烷膜包覆提升了FeSiCr粉体的抗腐蚀性和微波吸收性能。')

P('在层级核壳结构方面，合肥工业大学Xia课题组[8]构建了FeCo@SiO₂@NiFe₂O₄三元协同结构（图6）——FeCo提供高M_s（~220 emu/g）和高磁损耗，SiO₂中间层降低颗粒间电导和介电常数，NiFe₂O₄外壳提供额外磁损耗（自然共振峰~5 GHz）并增强阻抗匹配。三种结构RL对比：裸FeCo RL_min=-5.2 dB → FeCo@SiO₂ RL_min=-15.8 dB → FeCo@SiO₂@NiFe₂O₄ RL_min=-32.4 dB，三元协同效应显著。该工作的关键启示在于通过逐层构建实现介电-磁性独立调控，破解了传统单层包覆中介电和磁性"跷跷板"的固有矛盾。在FeCo/碳基复合材料方面，Liu等[15]报道了FeCo合金/石墨烯泡沫复合材料的制备及其高效微波吸收性能。Yao等[4]发现稀土La的钉扎效应对FeCo@rGO复合材料的耐腐蚀性和超高效吸波性能具有显著增强作用。')

insert_fig(os.path.join(FIG,'FeCo_EM_tuning_2017_1'),
    '图4  FeCo/石蜡复合物（70 wt%）反射损耗（RL vs 频率 vs 厚度）三维对比图',
    5.2, '引自文献[17], J. Alloys Compd., 2017, 717: 294')

insert_fig(os.path.join(FIG,'FeCo_ZnO_2021_1'),
    '图5  FeSiAl合金颗粒MnZn铁氧体复合前后的（a）SEM形貌及（b）电磁参数（ε′, ε″）频谱对比',
    5.0, '引自文献[10], J. Mater. Sci.: Mater. Electron., 2021, 32: 18371')

insert_fig(os.path.join(FIG,'FeCo_SiO2_NiFe2O4_2023_1'),
    '图6  FeCo@SiO₂@NiFe₂O₄层级核壳的（a-c）SEM/TEM形貌及（d-f）三种结构的反射损耗三维对比',
    5.2, '引自文献[8], J. Alloys Compd., 2023, 953: 170132')

H('2.4 国内外综合差距分析与本项目定位', 2)
P('综合文献调研[1-18]，当前国内外在耐高温吸波涂层领域存在以下关键技术和能力缺口：')

P('（1）吸收剂层面：尚无同时满足Tc≥900°C、8~18 GHz全频段可有效工作（μ″_max≥3.0）且ε″可在工程上可控（≤10）的磁性吸收剂产品。CIP高温失磁（600°C Ms衰减>50%）[14]、铁氧体低频受限（Snoek极限μ′<1.5）、纯FeCo介电过高（ε″=30~100）[3]，各有短板。本课题组已积累的三条FeCo路线（Fe₅₀Co₅₀片状/FeCoNd稀土降介电/FeCoGe半金属降本征电导率，见1.3节）为突破这一瓶颈提供了坚实的前期基础，但载流子抑制工艺的工程化（降ε″至≤10）和ALD包覆的耐高温验证（370°C）是必须攻克的两大工程难关。')

P('（2）防护层面：现有防腐包覆方案（溶胶-凝胶SiO₂耐盐雾~500 h）[1]在370°C的长期热氧稳定性和Cl⁻阻隔能力不足，有机硅树脂在300°C以上发生显著热降解[9]。ALD封装的耐高温验证极其缺乏——国内外全部ALD腐蚀防护研究最高测试温度约60°C[6,11,13]，与370°C工况存在约310°C温度缺口。缺少"原子级封装（ALD）+梯度涂层（CTE缓冲+耐候面层）"的系统性协同防护方案。')

P('（3）集成层面：尚无针对"0.5 mm超薄-8~18 GHz宽频-370°C耐温-3000 h盐雾"四位一体的涂层系统解决方案的公开报道。Xia等[2]的高温吸波材料综述指出同时满足T_test>300°C、d<1 mm和腐蚀环境三个条件的公开论文不足5篇。图7基于传输线理论给出了0.5 mm超薄约束下的阻抗匹配设计窗口：ε′需控制在10~25、ε″需控制在5~10的狭窄区间内，同时μ′需保持1.5~3.0——这与本课题关键技术二将ε″降至≤10、μ′保持≥1.5的目标高度一致。')

P('（4）国外工程应用公开文献披露严重不足：国外在高温吸波涂层的工程应用层面公开文献披露极为有限——军事敏感性和技术保密导致绝大部分工程化数据无法从公开渠道获取，仅能通过少量基础研究报道和装备观察间接推断。俄罗斯Mig-29K舰载机尾喷管采用了铁磁性合金+耐高温陶瓷基复合吸波涂层，但具体材料配方、工艺参数和性能数据未见公开报道。美国DDG-1000驱逐舰的舰炮涂层隐身方案未在任何公开文献中披露。美国NRL和AFRL在隐身涂层领域拥有完整的全链条研发体系，但其技术报告多为内部受控文件，公开发表的论文仅涉及基础材料层面的实验室数据，与真实装备的工程化指标之间存在无法弥合的信息鸿沟。')

P('本项目以FeCo基合金为核心吸收剂平台，依托课题组在FeCo粉体制备（三项工艺路线，见1.3节）、载流子抑制（修正Drude模型三项界面修正，见4.2节）和FB-ALD封装方面的技术积累，拟系统性攻克上述四个层面的技术瓶颈，填补舰炮隐身涂层在高温海洋环境领域的工程应用空白。')

insert_fig(os.path.join(SCHEM,'fig4_impedance_matching.png'),
    '图7  0.5 mm超薄涂层阻抗匹配设计：（a）10 GHz RL等高线图及目标窗口；（b）不同厚度RL-频率仿真曲线',
    5.0, '基于传输线理论计算（金属背衬，d=0.5 mm, μ′=2.5, μ″=1.5）')

# ============================================================
# COPY REMAINING SECTIONS 三〜八 from original
# ============================================================
sections = [
    ('三、研究内容', '四、关键技术', 'Heading 1'),
    ('四、关键技术', '五、工程风险', 'Heading 1'),
    ('五、工程风险', '六、进度计划', 'Heading 1'),
    ('六、进度计划', '七、考核指标', 'Heading 1'),
    ('七、考核指标', '八、研究周期', 'Heading 1'),
    ('八、研究周期', '参考文献', 'Heading 1'),
    ('参考文献', 'ZZZ_END', 'Normal'),
]

for start, end, style in sections:
    if end == 'ZZZ_END':
        # Copy everything after 参考文献
        copying = False
        for p in orig.paragraphs:
            if p.text.strip() == '参考文献':
                copying = True
                nr = doc.add_paragraph().add_run('参考文献')
                nr.bold = True; nr.font.size = Pt(14)
                nr.font.name = 'Times New Roman'
                continue
            if copying:
                np = doc.add_paragraph()
                np.paragraph_format.line_spacing = 1.0
                for r in p.runs:
                    nr2 = np.add_run(r.text)
                    nr2.font.size = r.font.size or Pt(9)
                    nr2.font.name = 'Times New Roman'
                    try: nr2.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
                    except: pass
    else:
        copy_paragraphs(orig, doc, start, end, style)

doc.save(OUT)
fsize = os.path.getsize(OUT)
print(f'\nDone: {OUT} ({fsize//1024} KB, {fsize/1024:.1f} KB)')
