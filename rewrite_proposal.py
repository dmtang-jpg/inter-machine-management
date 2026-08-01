#!/usr/bin/env python3
"""Replace 研究现状 section in the full proposal with expanded version."""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from lxml import etree
import os, re

FIG = os.path.expanduser('~/workspace/figures_literature/from_papers')
SCHEM = os.path.expanduser('~/workspace/figures_literature')
SRC = os.path.expanduser('~/workspace/耐高温磁性吸收剂_大项目申报_备份_v2.docx')
OUT = os.path.expanduser('~/workspace/耐高温磁性吸收剂_大项目申报_更新.docx')

doc = Document(SRC)

# Find the research status section boundaries
# Look for "二、国内外研究现状" heading and "三、研究内容" heading
section2_start = None
section3_start = None
paras_to_remove = []

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if '二、国内外研究现状' in text and p.style.name.startswith('Heading'):
        section2_start = i
    elif '三、研究内容' in text and p.style.name.startswith('Heading') and section2_start is not None:
        section3_start = i
        break

if section2_start is None or section3_start is None:
    print(f'ERROR: section2_start={section2_start}, section3_start={section3_start}')
    exit(1)

print(f'Found 研究现状 section: paragraphs {section2_start} to {section3_start-1}')

# Collect all paragraph elements to remove
elements_to_remove = []
for i in range(section2_start, section3_start):
    elements_to_remove.append(doc.paragraphs[i]._element)

# Remove old paragraphs
parent = doc.paragraphs[section2_start]._element.getparent()
for elem in elements_to_remove:
    parent.remove(elem)

# Now insert new content before the section 3 heading
# Find the new insertion point (where section 2 was)
insert_before = doc.paragraphs[section2_start]._element if section2_start < len(doc.paragraphs) else None
if insert_before is None:
    # Find section 3 heading element
    for p in doc.paragraphs:
        if '三、研究内容' in p.text and p.style.name.startswith('Heading'):
            insert_before = p._element
            break

print(f'Insert before: {insert_before.text[:50] if insert_before is not None else "NOT FOUND"}')

# Helper to create a paragraph with proper formatting
def make_para(doc, text, bold=False, indent=True, sz=12, heading_level=None):
    if heading_level:
        p = doc.add_paragraph()
        p.style = doc.styles[f'Heading {heading_level}']
        # Clear and add text
        p.clear()
        r = p.add_run(text)
        r.font.name = 'Times New Roman'
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体' if heading_level < 3 else '楷体')
        r.font.size = Pt(16 if heading_level == 1 else 14)
        r.bold = True
        r.font.color.rgb = RGBColor(0,0,0)
    else:
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.name = 'Times New Roman'
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        r.element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
        r.element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        r.font.size = Pt(sz)
        r.bold = bold
        if indent:
            p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(2)
    return p

def add_fig(p_after, path, caption, width=5.0, source=''):
    """Insert figure paragraph after given paragraph."""
    if not os.path.exists(path):
        base = os.path.splitext(path)[0]
        alt = f'{base}_clean.png'
        if os.path.exists(alt): path = alt
    if not os.path.exists(path):
        print(f'  WARNING: {path} not found')
        return
    
    # Create figure paragraph
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_img.add_run()
    r.add_picture(path, width=Inches(width))
    p_img.paragraph_format.space_before = Pt(6)
    p_img.paragraph_format.space_after = Pt(2)
    
    # Create caption paragraph
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p_cap.add_run(caption)
    r2.font.name = 'Times New Roman'; r2.font.size = Pt(9); r2.bold = True
    r2.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if source:
        r3 = p_cap.add_run(f'\n{source}')
        r3.font.size = Pt(7); r3.italic = True; r3.font.name = 'Times New Roman'
    p_cap.paragraph_format.space_after = Pt(8)
    
    # Move after target paragraph
    p_after._element.addnext(p_img._element)
    p_img._element.addnext(p_cap._element)
    return p_cap

def f(p_after, name, caption, width=5.0, source=''):
    return add_fig(p_after, os.path.join(FIG, name), caption, width, source)

# ============================================================
# Build new content
# ============================================================

# Section heading
h2 = make_para(doc, '二、国内外研究现状与技术差距', heading_level=1, bold=True)
insert_before.addprevious(h2._element)

# Helper: add para then return it for figure insertion
def P(text, bold=False, indent=True):
    p = make_para(doc, text, bold=bold, indent=indent)
    # Move before section 3
    insert_before.addprevious(p._element)
    return p

def H2(text):
    p = make_para(doc, text, heading_level=2)
    insert_before.addprevious(p._element)
    return p

# ============================================================
# 2.1
# ============================================================
H2('2.1 磁性金属吸收剂的防腐与耐温研究进展')

P('羰基铁粉（CIP）是当前工程应用最广泛的磁性吸波填料，占据隐身涂料市场70%以上份额。但其在高温海洋环境中的两个根本缺陷限制了舰炮工况的应用：（1）"居里温度陷阱"——CIP标称Tc≈770°C，但在600°C以上Ms衰减已超50%，μ′从常温3~4降至1.2~1.5，实质上丧失了微波磁损耗能力。这是因为多晶CIP中晶界、位错和杂质导致"有效居里温度"远低于体相值。（2）"盐雾腐蚀薄弱"——CIP在Cl⁻环境中Fe→Fe²⁺+2e⁻生成无磁性α-Fe₂O₃/FeOOH。')

P('在CIP表面改性方面，Li等采用ALD技术在CIP表面沉积Al₂O₃薄膜，将腐蚀电流密度从10⁻⁵降至10⁻⁷ A/cm²（降低2个数量级），8~12 GHz ε″下降约25%。图1展示了ALD Al₂O₃包覆CIP的SEM截面、XRD物相、XPS化学态和TGA热稳定性综合表征，从微观结构到宏观性能系统验证了ALD涂层在CIP防腐和电磁调控中的有效性。然而，这些研究均局限于200°C以下防腐验证，ALD涂层在370°C高温热循环下的界面应力退化行为几乎为空白领域——这正是本课题在关键技术一中的核心突破方向。')

last_p = P('在FeCo基吸收剂方面，Zare等通过理论计算指出片状FeCo径厚比>30时ε″急剧攀升至>100，形成"高磁导率-高介电"的两难困境（图2）。Gao等通过水热法在FeCo纳米颗粒表面外延生长ZnO半导体壳层（禁带宽度3.37 eV），利用ZnO高电阻率阻隔颗粒间电导，将ε″从60降至约18，但ZnO壳层在370°C下氧空位浓度改变可能导致电阻率下降。本课题组前期对FeCo合金的三条技术路线（等原子比Fe₅₀Co₅₀、稀土Nd掺杂、半金属Ge掺杂）已进行了系统实验积累：Fe₅₀Co₅₀片状粉体Tc=970°C；FeCoNd（x=5）在2 GHz处μ′=5.5, μ″_max=4.3, ε″≈8，实现了高磁损耗与可控介电的最佳折中；FeCoGe（x=0.2）1.5 mm厚度时RL=-10.80 dB。上述工作验证了FeCo基合金在高温吸波领域的可行性，为本项目降介电-超薄化提供了坚实的前期基础。')

f(last_p, 'CIP_ALD_anticorrosion_2022_1_clean.png',
    '图1  ALD Al₂O₃涂层羰基铁粉的（a）SEM/TEM形貌、（b）XRD图谱、（c-d）XPS Al 2p/O 1s能谱及（e）TGA热分析曲线',
    5.0, 'Li et al., Surf. Coat. Technol., 2022, 437: 128346')

last_p2 = P('', indent=False)  # spacer
f(last_p2, 'Zare_FeCo_2020_3_clean.png',
    '图2  不同还原工艺条件下FeCo合金颗粒的形貌演变SEM：（a）0.1 M→球形；（b）0.5 M→不规则团聚；（c-d）1.0-2.0 M→片状',
    4.8, 'Zare & Rhee, Mater. Res. Express, 2020, 7: 036516')

# ============================================================
# 2.2
# ============================================================
H2('2.2 原子层沉积（ALD）粉体封装技术的进展与不足')

P('ALD技术基于自限制表面化学反应，具有原子级厚度控制精度（Al₂O₃ GPC≈0.11 nm/cycle）、无针孔致密性（100%表面覆盖 vs 溶胶-凝胶5~15%残余孔隙率）和360°保形覆盖（TEM截面显示棱角处膜厚偏差<15%）三大独特优势。流化床ALD（FB-ALD）是粉体ALD的工程实现形式。Zhang等设计机械振动辅助流化床，在铁粉表面沉积5~10 nm Al₂O₃薄膜，将i_corr降至10⁻⁷ A/cm²，同时8~12 GHz ε″下降约25%。Gong等优化了FB-ALD前驱体脉冲/吹扫/流化参数，单批次处理量提升至200 g。Marin等对ALD Al₂O₃涂层不锈钢的长期（>1000 h）耐腐蚀性能进行了系统评估——这是ALD腐蚀防护领域的里程碑工作。')

last_p = P('然而，目前国内外FB-ALD封装研究几乎全部局限于200°C以下防腐验证。ALD Al₂O₃薄膜在370°C高温热循环下的应力退化行为——Al₂O₃（CTE 7.4×10⁻⁶/K）与FeCo金属（CTE 12×10⁻⁶/K）热膨胀失配导致的膜层龟裂和界面剥离（ΔT=350°C时σ_th≈4.8 MPa）——几乎为研究空白。这构成本课题"关键技术一"的核心创新方向。')

f(last_p, 'ALD_Al2O3_corrosion_2011_1_clean.png',
    '图3  ALD Al₂O₃涂层不锈钢的（a）截面TEM形貌及（b）长期电化学阻抗谱（Bode图，|Z| vs 频率）演变',
    4.8, 'Marin et al., Mater. Corros., 2015, 66: 907-914')

# ============================================================
# 2.3
# ============================================================
H2('2.3 电磁参量调控与吸波性能优化的研究进展')

P('在FeCo合金电磁参量的定量调控方面，Li等通过化学还原工艺参数矩阵实验，系统研究了FeCo复合物复介电常数和复磁导率的调控规律（图4）。工艺优化后样品在1.5 mm厚度下8~12 GHz频段RL≤-10 dB带宽达3.2 GHz，RL_min=-28.5 dB。该工作建立了"工艺参数→电磁参量→吸波性能"的完整定量映射，其方法论对本课题的RSM工艺优化具有直接借鉴价值。')

last_p = P('在层级核壳结构的协同吸波设计方面，Xia等构建了FeCo@SiO₂@NiFe₂O₄三元协同结构——FeCo提供高磁损耗，SiO₂中间层降低介电常数，NiFe₂O₄外壳提供额外磁损耗并增强阻抗匹配。三种结构RL对比（图5）为：裸FeCo RL_min=-5.2 dB → FeCo@SiO₂ RL_min=-15.8 dB → FeCo@SiO₂@NiFe₂O₄ RL_min=-32.4 dB，三元协同效应显著。该工作的关键启示在于：通过逐层构建实现介电-磁性的独立调控——SiO₂层降介电、NiFe₂O₄层增磁损耗，破解了传统单层包覆中介电和磁性"跷跷板"的固有矛盾。这一策略与本课题组"方案A可控氧化降介电+FB-ALD精密封装+三层梯度防护"的集成思路高度吻合。')

f(last_p, 'FeCo_EM_tuning_2017_1_clean.png',
    '图4  不同工艺参数制备的FeCo/石蜡复合物（70 wt%）反射损耗（RL vs 频率 vs 厚度）三维对比图',
    5.2, 'Li et al., J. Alloys Compd., 2017, 717: 294-302')

f(last_p, 'FeCo_SiO2_NiFe2O4_2023_1_clean.png',
    '图5  FeCo@SiO₂@NiFe₂O₄层级核壳结构的（a-c）SEM/TEM形貌及（d-f）三种结构反射损耗三维对比',
    5.2, 'Xia et al., J. Alloys Compd., 2023, 953: 170132')

# ============================================================
# 2.4
# ============================================================
H2('2.4 国内外综合差距分析与本项目定位')

P('综合文献调研，当前国内外在耐高温吸波涂层领域存在以下关键技术和能力缺口：')

P('（1）吸收剂层面：尚无同时满足Tc≥900°C、8~18 GHz全频段可有效工作（μ″_max≥3.0）且ε″可在工程上可控（≤10）的磁性吸收剂产品。CIP高温失磁（600°C Ms衰减>50%）、铁氧体低频受限（Snoek极限μ′<1.5）、纯FeCo介电过高（ε″=30~100），各有短板。本课题组已积累的三条FeCo路线（Fe₅₀Co₅₀片状/FeCoNd稀土降介电/FeCoGe半金属降本征电导率）为突破这一瓶颈提供了坚实的前期基础，但载流子抑制工艺的工程化（降ε″至≤10）和ALD包覆的耐高温验证（370°C）是必须攻克的两大工程难关。')

P('（2）防护层面：现有防腐包覆方案（溶胶-凝胶SiO₂耐盐雾~500 h）在370°C的长期热氧稳定性和Cl⁻阻隔能力不足。ALD封装的耐高温验证极其缺乏——国内外全部FB-ALD研究最高测试温度200°C，与370°C工况存在约170°C温度缺口。缺少"原子级封装（ALD）+梯度涂层（有机硅CTE缓冲+氟硅耐候）"的系统性协同防护方案。')

P('（3）集成层面：尚无针对"0.5 mm超薄-8~18 GHz宽频-370°C耐温-3000 h盐雾"四位一体的涂层系统解决方案的公开报道。Liu等的综述指出同时满足T_test>300°C、d<1 mm和腐蚀环境三个条件的公开论文不足5篇，印证了该方向的巨大创新空间和工程价值。')

P('本项目以FeCo基合金为核心吸收剂平台，依托课题组在FeCo粉体制备（三项工艺路线）、载流子抑制（修正Drude模型三项界面修正）和FB-ALD封装方面的技术积累，拟系统性攻克上述三个层面的技术瓶颈，填补舰炮隐身涂层在高温海洋环境领域的工程应用空白。')

# ============================================================
# Save
# ============================================================
doc.save(OUT)
fsize = os.path.getsize(OUT)
print(f'\nDone: {OUT} ({fsize//1024} KB, {fsize/1024:.0f} KB)')
