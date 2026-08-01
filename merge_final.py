#!/usr/bin/env python3
"""Merge standalone review section 二 back into full proposal, then save."""
from docx import Document
from docx.oxml.ns import qn
from lxml import etree
import copy, os

# Load source documents
standalone = Document(os.path.expanduser('~/workspace/国内外研究现状_独立版.docx'))
full = Document(os.path.expanduser('~/workspace/耐高温磁性吸收剂_大项目申报_备份_v2.docx'))
OUT = os.path.expanduser('~/workspace/耐高温磁性吸收剂_大项目申报_更新.docx')

# === STEP 1: Remove old section 二 from full proposal ===
section2_start = None
section3_start = None
for i, p in enumerate(full.paragraphs):
    t = p.text.strip()
    if '二、国内外研究现状' in t and p.style.name.startswith('Heading'):
        section2_start = i
    elif '三、研究内容' in t and p.style.name.startswith('Heading') and section2_start is not None:
        section3_start = i
        break

print(f'Old section 二: paras {section2_start}-{section3_start-1}')

# Remove old paragraphs
body = full.paragraphs[section2_start]._element.getparent()
elements_to_remove = []
for i in range(section2_start, section3_start):
    elements_to_remove.append(full.paragraphs[i]._element)
for elem in elements_to_remove:
    body.remove(elem)

# === STEP 2: Extract new section content from standalone ===
# Find the body element
standalone_body = standalone.element.body

# Find section 二 heading in standalone
standalone_s2_heading = None
for p in standalone.paragraphs:
    if '一、国内研究现状' in p.text and p.style.name.startswith('Heading'):
        standalone_s2_heading = p._element
        break

# Collect all elements from this heading to the end (before references)
elements_to_copy = []
started = False
for child in standalone_body:
    if child is standalone_s2_heading:
        started = True
    if started:
        # Stop before 参考文献
        if child.tag.endswith('}p'):
            # Check if this paragraph contains 参考文献 as heading
            for run in child.findall('.//' + qn('w:r')):
                texts = [t.text or '' for t in run.findall(qn('w:t'))]
                if '参考文献' in ''.join(texts):
                    # Check if it's a heading style
                    pPr = child.find(qn('w:pPr'))
                    if pPr is not None:
                        pStyle = pPr.find(qn('w:pStyle'))
                        if pStyle is not None and 'Heading' in (pStyle.get(qn('w:val')) or ''):
                            started = False
                            break
        if started:
            elements_to_copy.append(child)

print(f'Copied {len(elements_to_copy)} elements from standalone')

# === STEP 3: Insert before section 三 ===
# Find section 三 heading in full document
section3_elem = None
for p in full.paragraphs:
    if '三、研究内容' in p.text and p.style.name.startswith('Heading'):
        section3_elem = p._element
        break

if section3_elem is None:
    print('ERROR: section 三 not found')
    exit(1)

# Insert each element before section 三 (reverse order to maintain sequence)
for elem in reversed(elements_to_copy):
    section3_elem.addprevious(copy.deepcopy(elem))

# === STEP 4: Save ===
full.save(OUT)
fsize = os.path.getsize(OUT)
print(f'\nDone: {OUT} ({fsize//1024} KB)')

# Quick verify
verify = Document(OUT)
h_count = sum(1 for p in verify.paragraphs if p.style.name.startswith('Heading'))
img_count = sum(1 for rel in verify.part.rels.values() if 'image' in rel.reltype)
print(f'Headings: {h_count}, Images: {img_count}')
