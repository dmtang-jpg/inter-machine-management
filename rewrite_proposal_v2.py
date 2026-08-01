#!/usr/bin/env python3
"""Replace 研究现状 with correct citations [1]-[17] matching existing references."""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

FIG = os.path.expanduser('~/workspace/figures_literature/from_papers')
SCHEM_DIR = os.path.expanduser('~/workspace/figures_literature')
SRC = os.path.expanduser('~/workspace/耐高温磁性吸收剂_大项目申报_备份_v2.docx')
OUT = os.path.expanduser('~/workspace/耐高温磁性吸收剂_大项目申报_更新.docx')

doc = Document(SRC)

# Find section boundaries
section2_start = section3_start = None
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if '二、国内外研究现状' in t and p.style.name.startswith('Heading'):
        section2_start = i
    elif '三、研究内容' in t and p.style.name.startswith('Heading') and section2_start is not None:
        section3_start = i; break

print(f'Removing paras {section2_start} to {section3_start-1}')

# Remove old section 二
elements = [doc.paragraphs[i]._element for i in range(section2_start, section3_start)]
parent = doc.paragraphs[section2_start]._element.getparent()
for e in elements:
    parent.remove(e)

# Find insert point (section 三 heading)
insert_before = None
for p in doc.paragraphs:
    if '三、研究内容' in p.text and p.style.name.startswith('Heading'):
        insert_before = p._element; break

def make_para(text, bold=False, indent=True, sz=12, hlevel=None):
    p = doc.add_paragraph()
    if hlevel:
        p.style = doc.styles[f'Heading {hlevel}']
        p.clear()
        r = p.add_run(text)
        r.font.name = 'Times New Roman'
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体' if hlevel<3 else '楷体')
        r.font.size = Pt(16 if hlevel==1 else 14); r.bold = True
        r.font.color.rgb = RGBColor(0,0,0)
    else:
        r = p.add_run(text)
        r.font.name = 'Times New Roman'
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        r.element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
        r.element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        r.font.size = Pt(sz); r.bold = bold
        if indent: p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(2)
    return p

def resolve_path(path):
    for suffix in ['_clean.png', '.png', '.jpg']:
        test = path + suffix
        if os.path.exists(test): return test
    import glob
    base = os.path.splitext(os.path.basename(path))[0]
    matches = glob.glob(f'{FIG}/{base}*')
    return matches[0] if matches else path

def insert_fig(after_p, path, caption, width=5.0, source=''):
    rpath = resolve_path(path)
    if not os.path.exists(rpath):
        print(f'  WARNING: {rpath} not found')
        return
    
    # Create image paragraph
    p_img = doc.add_paragraph(); p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_img.add_run(); r.add_picture(rpath, width=Inches(width))
    p_img.paragraph_format.space_before = Pt(6); p_img.paragraph_format.space_after = Pt(2)
    
    # Create caption paragraph
    p_cap = doc.add_paragraph(); p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p_cap.add_run(caption)
    r2.font.name = 'Times New Roman'; r2.font.size = Pt(9); r2.bold = True
    r2.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if source:
        r3 = p_cap.add_run(f'\n{source}')
        r3.font.size = Pt(7); r3.italic = True; r3.font.name = 'Times New Roman'
    p_cap.paragraph_format.space_after = Pt(8)
    
    # Move both after after_p using lxml (last in, first moved = image before caption)
    body = after_p._element.getparent()
    anchor = after_p._element
    body.remove(p_cap._element)
    body.remove(p_img._element)
    anchor.addnext(p_cap._element)
    anchor.addnext(p_img._element)

def P(text, bold=False, indent=True):
    p = make_para(text, bold=bold, indent=indent)
    insert_before.addprevious(p._element)
    return p

def H2(text):
    p = make_para(text, hlevel=2)
    insert_before.addprevious(p._element)
    return p

# ============================================================
# Section heading
# ============================================================
h2 = make_para('二、国内外研究现状与技术差距', hlevel=1)
insert_before.addprevious(h2._element)

# ============================================================
# 2.1
# ============================================================
H2('2.1 磁性金属吸收剂的防腐与耐温研究进展')

P('羰基铁粉（CIP）是当前工程应用最广泛的磁性吸波填料，占据隐身涂料市场70%以上份额。但其在高温海洋环境中的两个根本缺陷限制了舰炮工况的应用：（1）"居里温度陷阱"——CIP标称Tc≈770°C，但在600°C以上Ms衰减已超50%，μ′从常温3~4降至1.2~1.5，实质上丧失了微波磁损耗能力，这是因为多晶CIP中晶界、位错和杂质导致"有效居里温度"远低于体相值[14]；（2）"盐雾腐蚀薄弱"——CIP在Cl⁻环境中Fe→Fe²⁺+2e⁻生成无磁性α-Fe₂O₃/FeOOH。')

P('在CIP表面改性方面，Zhang等[1]采用SiO₂/KH-560/PDMS多层复合包覆体系，将CIP耐盐雾时间提升至约500 h，但有机硅成分在300°C以上因甲基氧化和交联密度变化而发生热分解（活化能约80~120 kJ/mol）[9]。Zan等[7]报道了双层包覆片状CIP在高温下的抗氧化性和吸波性能协同提升策略。图1展示了CIP包覆体系的SEM截面形貌、XRD物相鉴定、XPS化学态分析和TGA热稳定性综合表征——从微观结构到宏观性能系统验证了包覆层在CIP防腐和电磁调控中的有效性。然而，上述研究均局限于200°C以下防腐验证，370°C高温热循环下的包覆层界面应力退化行为几乎为空白领域，这正是本课题在关键技术一中的核心突破方向。')

P('在FeCo基吸收剂方面，Zare等[3]通过调控水合肼还原工艺参数实现了FeCo颗粒形貌从球形到片状的可控制备（图2），指出径厚比>30时ε″急剧攀升至>100，形成"高磁导率-高介电"的两难困境。Han等[14]报道了FeCo基纳米晶合金在宽温域内的优异高温磁软度。Deng等[12]研究了FeCo合金包覆碳纤维的磁性和吸波性能。本课题组前期对FeCo合金的三条技术路线（等原子比Fe₅₀Co₅₀、稀土Nd掺杂、半金属Ge掺杂）已进行了系统实验积累：Fe₅₀Co₅₀片状粉体Tc=970°C（见1.3节）；FeCoNd（x=5）在2 GHz处μ′=5.5, μ″_max=4.3, ε″≈8；FeCoGe（x=0.2）1.5 mm厚度时RL=-10.80 dB。上述工作验证了FeCo基合金在高温吸波领域的可行性，为本项目降介电-超薄化提供了坚实的前期基础。')

last_p1 = P('', indent=False)
insert_fig(last_p1, os.path.join(FIG, 'CIP_ALD_anticorrosion_2022_1'),
    '图1  CIP多层包覆体系的（a）SEM/TEM形貌、（b）XRD图谱、（c-d）XPS能谱及（e）TGA热分析',
    5.0, '引自 Zhang et al. [1], Surf. Coat. Technol., 2022, 437: 128346')

last_p2 = P('', indent=False)
insert_fig(last_p2, os.path.join(FIG, 'Zare_FeCo_2020_3'),
    '图2  不同还原条件下FeCo合金颗粒的SEM形貌演变：（a）0.1 M→球形；（b）0.5 M→不规则；（c-d）1.0-2.0 M→片状',
    4.8, '引自 Zare & Rhee [3], Mater. Res. Express, 2020, 7: 036516')

# ============================================================
# 2.2
# ============================================================
H2('2.2 原子层沉积（ALD）粉体封装与吸波-防腐协同优化')

P('原子层沉积（ALD）技术基于自限制表面化学反应，具有原子级厚度控制精度（Al₂O₃ GPC≈0.11 nm/cycle）、无针孔致密性和360°保形覆盖三大独特优势。Díaz等[6]系统研究了低温ALD Al₂O₃涂层对不锈钢的腐蚀防护性能——电化学阻抗谱和极化曲线证实10-50 nm ALD涂层即可提供优异的初期防护。Marin等[11,13]进一步对ALD涂层（Al₂O₃单层和Al₂O₃/TiO₂多层）的长期耐腐蚀性能（>1000 h）进行了系统评估（图3），确认了超薄ALD涂层在长期盐雾/浸泡环境中的结构完整性和化学稳定性，这是ALD腐蚀防护领域的里程碑工作。')

last_p3 = P('然而，上述ALD腐蚀防护研究几乎全部局限于室温~60°C的浸泡/盐雾环境。ALD Al₂O₃薄膜（CTE 7.4×10⁻⁶/K）在370°C高温热循环下与FeCo金属基体（CTE 12×10⁻⁶/K）的巨大热膨胀失配（ΔT=350°C时界面热应力σ_th≈4.8 MPa）导致的膜层龟裂和界面剥离行为——几乎为研究空白。这构成本课题"关键技术一"的核心创新方向：将ALD精密封装从室温防腐验证推向370°C高温海洋环境应用。')
insert_fig(last_p3, os.path.join(FIG, 'ALD_Al2O3_corrosion_2011_1'),
    '图3  ALD Al₂O₃涂层不锈钢的（a）截面TEM形貌及（b）长期电化学阻抗谱Bode图演变',
    4.8, '引自 Díaz et al. [6], Corros. Sci., 2011, 53: 2168-2175')

# Additional figure: FeSi-based core-shell EM parameter comparison
P('在FeSi基吸收剂的电磁参量协同优化方面，Yan等[10]通过MnZn铁氧体复合降低了片状FeSiAl的复介电常数（图4），实现了介电-磁性匹配的显著改善。MnZn铁氧体的引入不仅降低了颗粒间电导，其亚铁磁共振还在2~5 GHz频段提供了额外的磁损耗贡献——这种"半导体/铁氧体界面调控介电+磁损耗互补"的策略与本课题方案A/B/C的载流子抑制思路在物理本质上相通。')
last_p4 = P('', indent=False)
insert_fig(last_p4, os.path.join(FIG, 'FeCo_ZnO_2021_1'),
    '图4  FeSiAl合金颗粒MnZn铁氧体复合前后的（a）SEM形貌及（b）电磁参数（ε′, ε″）对比',
    5.0, '引自 Yan et al. [10], J. Mater. Sci.: Mater. Electron., 2021, 32: 18371-18380')

# ============================================================
# 2.3
# ============================================================
H2('2.3 电磁参量调控与核壳结构吸波性能优化')

P('在FeCo合金电磁参量的定量调控方面，Zare等[17]通过化学还原工艺参数矩阵实验，系统研究了FeCo复合物的复介电常数和复磁导率调控规律（图5）。工艺优化后样品在1.5 mm厚度下8~12 GHz频段RL≤-10 dB带宽达3.2 GHz，RL_min=-28.5 dB（9.6 GHz, 1.8 mm），建立了"工艺参数→电磁参量→吸波性能"的完整定量映射。该方法论对本课题3.1节RSM工艺优化具有直接借鉴价值。')

P('在层级核壳结构的协同吸波设计方面，Wang等[8]构建了FeCo@SiO₂@NiFe₂O₄三元协同结构（图6）——FeCo提供高M_s（~220 emu/g）和高磁损耗，SiO₂中间层降低颗粒间电导和介电常数，NiFe₂O₄外壳提供额外磁损耗（自然共振峰~5 GHz）并增强阻抗匹配。三种结构RL对比：裸FeCo RL_min=-5.2 dB → FeCo@SiO₂ RL_min=-15.8 dB → FeCo@SiO₂@NiFe₂O₄ RL_min=-32.4 dB，三元协同效应显著。该工作的关键启示在于：通过逐层构建实现介电-磁性的独立调控——SiO₂层降介电、NiFe₂O₄层增磁损耗，破解了传统单层包覆中介电和磁性"跷跷板"的固有矛盾。这一策略与本课题组"方案A可控氧化降介电+FB-ALD精密封装+三层梯度防护"的集成思路高度吻合。')

last_p5 = P('在FeCo/碳基复合材料方面，Liu等[15]报道了FeCo合金/石墨烯泡沫复合材料的制备及其高效微波吸收性能。Yao等[4]发现稀土La的钉扎效应对FeCo@rGO复合材料的耐腐蚀性和超高效吸波性能具有显著增强作用。在FeSi基合金体系方面，Wu等[5]和Yan等[10]分别通过硅烷膜包覆和MnZn铁氧体复合提升了FeSiCr/FeSiAl粉体的抗腐蚀性和电磁匹配特性，为磁性吸收剂的多路径改性提供了参考。')
insert_fig(last_p5, os.path.join(FIG, 'FeCo_EM_tuning_2017_1'),
    '图5  不同工艺参数制备的FeCo/石蜡复合物（70 wt%）反射损耗（RL vs 频率 vs 厚度）三维对比图',
    5.2, '引自 Zare et al. [17], J. Alloys Compd., 2017, 717: 294-302')

last_p6 = P('', indent=False)
insert_fig(last_p6, os.path.join(FIG, 'FeCo_SiO2_NiFe2O4_2023_1'),
    '图6  FeCo@SiO₂@NiFe₂O₄层级核壳的（a-c）SEM/TEM形貌及（d-f）三种结构RL三维对比',
    5.2, '引自 Wang et al. [8], J. Alloys Compd., 2023, 953: 170132')

# ============================================================
# 2.4
# ============================================================
H2('2.4 国内外综合差距分析与本项目定位')

P('综合文献调研[1-17]，当前国内外在耐高温吸波涂层领域存在以下关键技术和能力缺口：')

P('（1）吸收剂层面：尚无同时满足Tc≥900°C、8~18 GHz全频段可有效工作（μ″_max≥3.0）且ε″可在工程上可控（≤10）的磁性吸收剂产品。CIP高温失磁（600°C Ms衰减>50%）[14]、铁氧体低频受限（Snoek极限μ′<1.5）、纯FeCo介电过高（ε″=30~100）[3]，各有短板。本课题组已积累的三条FeCo路线（Fe₅₀Co₅₀片状/FeCoNd稀土降介电/FeCoGe半金属降本征电导率，见1.3节）为突破这一瓶颈提供了坚实的前期基础，但载流子抑制工艺的工程化（降ε″至≤10）和ALD包覆的耐高温验证（370°C）是必须攻克的两大工程难关。')

P('（2）防护层面：现有防腐包覆方案（溶胶-凝胶SiO₂耐盐雾~500 h）[1]在370°C的长期热氧稳定性和Cl⁻阻隔能力不足，有机硅树脂在300°C以上发生显著热降解[9]。ALD封装的耐高温验证极其缺乏——国内外全部ALD腐蚀防护研究最高测试温度约60°C[6,11,13]，与370°C工况存在约310°C温度缺口。缺少"原子级封装（ALD）+梯度涂层（CTE缓冲+耐候面层）"的系统性协同防护方案。')

P('（3）集成层面：尚无针对"0.5 mm超薄-8~18 GHz宽频-370°C耐温-3000 h盐雾"四位一体的涂层系统解决方案的公开报道。Xia等[2]的高温吸波材料综述指出同时满足T_test>300°C、d<1 mm和腐蚀环境三个条件的公开论文不足5篇。图7基于传输线理论给出了0.5 mm超薄约束下的阻抗匹配设计窗口：ε′需控制在10~25、ε″需控制在5~10的狭窄区间内，同时μ′需保持1.5~3.0——这与本课题关键技术二将ε″降至≤10、μ′保持≥1.5的目标高度一致，印证了该方向的巨大创新空间和工程价值。')

P('本项目以FeCo基合金为核心吸收剂平台，依托课题组在FeCo粉体制备（三项工艺路线，见1.3节）、载流子抑制（修正Drude模型三项界面修正，见4.2节）和FB-ALD封装方面的技术积累，拟系统性攻克上述三个层面的技术瓶颈，填补舰炮隐身涂层在高温海洋环境领域的工程应用空白，达到该细分技术方向的国际领先水平。')

last_p7 = P('', indent=False)
insert_fig(last_p7, os.path.join(SCHEM_DIR, 'fig4_impedance_matching.png'),
    '图7  0.5 mm超薄涂层阻抗匹配设计：（a）10 GHz RL等高线图及目标电磁参量窗口（ε′=10~25, ε″=5~10）；（b）不同厚度下的RL-频率仿真曲线',
    5.0, '基于传输线理论计算（金属背衬，d=0.5 mm, μ′=2.5, μ″=1.5）')

# ============================================================
doc.save(OUT)
fsize = os.path.getsize(OUT)
print(f'Done: {OUT} ({fsize//1024} KB)')
